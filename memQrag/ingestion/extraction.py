"""Text extraction adapters for memQrag document ingestion (Phase 2 PR 2).

Consumes the `RawDocument` intake contract (see `memQrag.ingestion.contracts`)
and produces an `ExtractedDocument`: document-level metadata plus a list of
`ExtractedSegment` blocks carrying page number and section heading where the
source format makes that meaningful. Semantic chunking (Phase 2 PR 3) reads
this output; this module performs no chunking or persistence.

Adapter behavior per format, see docs/DECISIONS.md ("Text Extraction Adapter
Behavior"):

- PDF: one segment per page (`page_number` set, `section_heading` unset);
  `created_date`/`last_modified_date` come from the PDF Info dictionary when
  present.
- DOCX: one segment per section, split on paragraphs styled "Heading *"
  (`section_heading` set to the heading text, `page_number` unset, since
  DOCX has no fixed page boundaries independent of the rendering engine);
  dates come from the document's core properties.
- TXT: a single segment with no page number or heading; no embedded dates.
- Markdown: one segment per ATX (`#`) heading section; no embedded dates.
"""

from __future__ import annotations

import re
from collections.abc import Callable
from dataclasses import dataclass, field
from datetime import datetime
from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader

from memQrag.ingestion.contracts import RawDocument, SupportedFileType

_MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")


@dataclass(frozen=True)
class ExtractedSegment:
    """One extracted block of text with its structural metadata."""

    text: str
    page_number: int | None = None
    section_heading: str | None = None


@dataclass(frozen=True)
class ExtractedDocument:
    """The result of text extraction: document-level metadata plus segments."""

    source_document: str
    file_type: SupportedFileType
    created_date: datetime | None
    last_modified_date: datetime | None
    segments: list[ExtractedSegment] = field(default_factory=list)


def extract_text(document: RawDocument) -> ExtractedDocument:
    """Dispatch to the adapter matching `document.file_type` and extract text."""
    return _EXTRACTORS[document.file_type](document)


def _extract_txt(document: RawDocument) -> ExtractedDocument:
    text = document.content.decode("utf-8", errors="replace")
    segments = [ExtractedSegment(text=text)] if text.strip() else []
    return ExtractedDocument(
        source_document=document.filename,
        file_type=document.file_type,
        created_date=None,
        last_modified_date=None,
        segments=segments,
    )


def _extract_markdown(document: RawDocument) -> ExtractedDocument:
    text = document.content.decode("utf-8", errors="replace")
    segments: list[ExtractedSegment] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        block = "\n".join(current_lines).strip()
        if block:
            segments.append(ExtractedSegment(text=block, section_heading=current_heading))

    for line in text.splitlines():
        match = _MARKDOWN_HEADING_RE.match(line)
        if match:
            flush()
            current_heading = match.group(2).strip()
            current_lines = []
        else:
            current_lines.append(line)
    flush()

    return ExtractedDocument(
        source_document=document.filename,
        file_type=document.file_type,
        created_date=None,
        last_modified_date=None,
        segments=segments,
    )


def _extract_docx(document: RawDocument) -> ExtractedDocument:
    docx_document = DocxDocument(BytesIO(document.content))

    segments: list[ExtractedSegment] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        block = "\n".join(current_lines).strip()
        if block:
            segments.append(ExtractedSegment(text=block, section_heading=current_heading))

    for paragraph in docx_document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            flush()
            current_heading = paragraph.text.strip()
            current_lines = []
        elif paragraph.text.strip():
            current_lines.append(paragraph.text)
    flush()

    core_properties = docx_document.core_properties
    return ExtractedDocument(
        source_document=document.filename,
        file_type=document.file_type,
        created_date=core_properties.created,
        last_modified_date=core_properties.modified,
        segments=segments,
    )


def _extract_pdf(document: RawDocument) -> ExtractedDocument:
    reader = PdfReader(BytesIO(document.content))

    segments = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            segments.append(ExtractedSegment(text=page_text, page_number=page_number))

    info = reader.metadata
    created_date = info.creation_date if info else None
    last_modified_date = info.modification_date if info else None

    return ExtractedDocument(
        source_document=document.filename,
        file_type=document.file_type,
        created_date=created_date,
        last_modified_date=last_modified_date,
        segments=segments,
    )


_EXTRACTORS: dict[SupportedFileType, Callable[[RawDocument], ExtractedDocument]] = {
    SupportedFileType.PDF: _extract_pdf,
    SupportedFileType.DOCX: _extract_docx,
    SupportedFileType.TXT: _extract_txt,
    SupportedFileType.MARKDOWN: _extract_markdown,
}
